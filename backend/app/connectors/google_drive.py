import io
import json
from datetime import datetime

from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload

from app.connectors.base import BaseConnector, ConnectorStatus
from app.processing.pipeline import RawDocument

SCOPES_DRIVE = ["https://www.googleapis.com/auth/drive.readonly"]
SKIP_MIME = {
    "application/vnd.google-apps.folder",
    "video/",
    "audio/",
    "image/",
}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB


class GoogleDriveConnector(BaseConnector):
    async def authenticate(self, **kwargs) -> dict:
        return {}

    def _build_service(self, token_dict: dict):
        creds = Credentials(
            token=token_dict.get("access_token"),
            refresh_token=token_dict.get("refresh_token"),
            token_uri="https://oauth2.googleapis.com/token",
            client_id=token_dict.get("client_id"),
            client_secret=token_dict.get("client_secret"),
            scopes=SCOPES_DRIVE,
        )
        return build("drive", "v3", credentials=creds)

    async def fetch_documents(self, connector_record, since: datetime | None = None) -> list[RawDocument]:
        from app.connectors.utils import decrypt_token
        token_dict = decrypt_token(connector_record.oauth_token_encrypted)
        service = self._build_service(token_dict)

        query = "trashed=false and mimeType != 'application/vnd.google-apps.folder'"
        if since:
            since_str = since.strftime("%Y-%m-%dT%H:%M:%S")
            query += f" and modifiedTime > '{since_str}'"

        results = []
        page_token = None

        while True:
            kwargs = {
                "q": query,
                "fields": "nextPageToken, files(id, name, mimeType, modifiedTime, owners, webViewLink, size)",
                "pageSize": 100,
            }
            if page_token:
                kwargs["pageToken"] = page_token

            response = service.files().list(**kwargs).execute()
            files = response.get("files", [])

            for file in files:
                size = int(file.get("size", 0) or 0)
                mime = file.get("mimeType", "")

                if any(skip in mime for skip in SKIP_MIME):
                    continue
                if size > MAX_FILE_SIZE:
                    continue

                text = self._extract_text(service, file)
                if not text:
                    continue

                owner = file.get("owners", [{}])[0].get("emailAddress", "")
                modified = file.get("modifiedTime")
                dt = datetime.fromisoformat(modified.replace("Z", "+00:00")) if modified else None

                results.append(
                    RawDocument(
                        source_id=file["id"],
                        title=file.get("name", "Untitled"),
                        text=text,
                        author=owner,
                        source_url=file.get("webViewLink"),
                        source_created_at=dt,
                        is_html=False,
                    )
                )

            page_token = response.get("nextPageToken")
            if not page_token:
                break

        return results

    def _extract_text(self, service, file: dict) -> str:
        mime = file.get("mimeType", "")
        fid = file["id"]

        try:
            if mime == "application/vnd.google-apps.document":
                data = service.files().export(fileId=fid, mimeType="text/plain").execute()
                return data.decode("utf-8", errors="ignore")

            elif mime == "application/vnd.google-apps.spreadsheet":
                data = service.files().export(fileId=fid, mimeType="text/csv").execute()
                return data.decode("utf-8", errors="ignore")

            elif mime in ("text/plain", "text/markdown", "text/csv"):
                fh = io.BytesIO()
                request = service.files().get_media(fileId=fid)
                downloader = MediaIoBaseDownload(fh, request)
                done = False
                while not done:
                    _, done = downloader.next_chunk()
                return fh.getvalue().decode("utf-8", errors="ignore")

            elif mime == "application/pdf":
                fh = io.BytesIO()
                request = service.files().get_media(fileId=fid)
                downloader = MediaIoBaseDownload(fh, request)
                done = False
                while not done:
                    _, done = downloader.next_chunk()
                return _extract_pdf_text(fh.getvalue())

            elif mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                fh = io.BytesIO()
                request = service.files().get_media(fileId=fid)
                downloader = MediaIoBaseDownload(fh, request)
                done = False
                while not done:
                    _, done = downloader.next_chunk()
                return _extract_docx_text(fh.getvalue())

        except Exception:
            return ""

        return ""

    async def get_status(self, connector_record) -> ConnectorStatus:
        return ConnectorStatus(
            status=connector_record.status,
            last_sync_at=connector_record.last_sync_at,
            document_count=connector_record.document_count,
            error_message=connector_record.error_message,
        )


def _extract_pdf_text(data: bytes) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        texts = []
        for page in reader.pages:
            texts.append(page.extract_text() or "")
        return "\n".join(texts)
    except Exception:
        return ""


def _extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception:
        return ""
